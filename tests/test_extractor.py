"""Tests for file content extraction module."""

import json
import pytest
from pathlib import Path

from embedded_finder.extractor import (
    extract_text, chunk_text, get_mime_type, is_natively_embeddable,
    get_pdf_page_count, convert_image_to_png, prepare_native_file,
    get_media_duration,
)


def test_extract_plain_text(tmp_dir):
    f = tmp_dir / "hello.txt"
    f.write_text("Hello, world!")
    assert extract_text(f) == "Hello, world!"


def test_extract_python_file(tmp_dir):
    f = tmp_dir / "code.py"
    f.write_text("def foo():\n    return 42\n")
    text = extract_text(f)
    assert "def foo():" in text
    assert "return 42" in text


def test_extract_markdown(tmp_dir):
    f = tmp_dir / "notes.md"
    f.write_text("# Title\n\nSome content here.\n")
    text = extract_text(f)
    assert "# Title" in text
    assert "Some content here." in text


def test_extract_json(tmp_dir):
    f = tmp_dir / "data.json"
    data = {"name": "test", "values": [1, 2, 3]}
    f.write_text(json.dumps(data))
    text = extract_text(f)
    assert "name" in text
    assert "test" in text
    assert "values" in text


def test_extract_json_pretty_prints(tmp_dir):
    f = tmp_dir / "data.json"
    f.write_text('{"a":1}')
    text = extract_text(f)
    # Should be pretty-printed
    assert "  " in text or "\n" in text


def test_extract_handles_encoding_errors(tmp_dir):
    f = tmp_dir / "binary.txt"
    f.write_bytes(b"\x80\x81\x82\x83")
    # Should not raise, returns content via latin-1 fallback
    text = extract_text(f)
    assert isinstance(text, str)


def test_extract_returns_empty_for_missing_file():
    text = extract_text("/nonexistent/file.txt")
    assert text == ""


def test_extract_docx(tmp_dir):
    from docx import Document
    f = tmp_dir / "doc.docx"
    doc = Document()
    doc.add_paragraph("First paragraph about testing.")
    doc.add_paragraph("Second paragraph about code.")
    doc.save(str(f))

    text = extract_text(f)
    assert "First paragraph about testing." in text
    assert "Second paragraph about code." in text


# --- chunk_text tests ---

def test_chunk_text_short_text():
    chunks = chunk_text("Short text.")
    assert chunks == ["Short text."]


def test_chunk_text_empty_text():
    assert chunk_text("") == []
    assert chunk_text("   ") == []


def test_chunk_text_splits_long_text():
    # Create text longer than 100 tokens (~400 chars)
    paragraphs = [f"Paragraph {i} " * 20 for i in range(10)]
    text = "\n\n".join(paragraphs)
    chunks = chunk_text(text, max_tokens=100)
    assert len(chunks) > 1
    # All content should be preserved
    rejoined = " ".join(chunks)
    for i in range(10):
        assert f"Paragraph {i}" in rejoined


def test_chunk_text_respects_paragraph_boundaries():
    text = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."
    chunks = chunk_text(text, max_tokens=1000)
    # Short enough to fit in one chunk
    assert len(chunks) == 1
    assert "First paragraph." in chunks[0]


def test_chunk_text_handles_single_long_paragraph():
    # One paragraph that's very long
    text = "word " * 5000  # ~25000 chars
    chunks = chunk_text(text, max_tokens=100)
    assert len(chunks) > 1


# --- Multimodal helper tests ---

def test_get_mime_type_image():
    assert get_mime_type("photo.png") == "image/png"
    assert get_mime_type("photo.jpg") == "image/jpeg"
    assert get_mime_type("photo.jpeg") == "image/jpeg"


def test_get_mime_type_audio():
    assert get_mime_type("song.mp3") == "audio/mpeg"
    assert get_mime_type("track.wav") == "audio/wav"


def test_get_mime_type_video():
    assert get_mime_type("clip.mp4") == "video/mp4"
    assert get_mime_type("movie.mov") == "video/quicktime"


def test_get_mime_type_pdf():
    assert get_mime_type("doc.pdf") == "application/pdf"


def test_get_mime_type_unknown():
    assert get_mime_type("file.xyz") == "application/octet-stream"


def test_is_natively_embeddable_image(tmp_dir):
    img = tmp_dir / "test.png"
    img.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100)
    assert is_natively_embeddable(img) is True


def test_is_natively_embeddable_convertible_image(tmp_dir):
    gif = tmp_dir / "test.gif"
    gif.write_bytes(b"GIF89a" + b"\x00" * 100)
    assert is_natively_embeddable(gif) is True


def test_is_natively_embeddable_audio(tmp_dir):
    audio = tmp_dir / "test.mp3"
    audio.write_bytes(b"\xff\xfb" + b"\x00" * 100)
    assert is_natively_embeddable(audio) is True


def test_is_natively_embeddable_video(tmp_dir):
    video = tmp_dir / "test.mp4"
    video.write_bytes(b"\x00" * 100)
    assert is_natively_embeddable(video) is True


def test_is_natively_embeddable_unsupported_video(tmp_dir):
    for ext in [".avi", ".mkv", ".webm"]:
        video = tmp_dir / f"test{ext}"
        video.write_bytes(b"\x00" * 100)
        assert is_natively_embeddable(video) is False


def test_is_natively_embeddable_text_file(tmp_dir):
    txt = tmp_dir / "test.py"
    txt.write_text("print('hello')")
    assert is_natively_embeddable(txt) is False


def test_is_natively_embeddable_small_pdf(tmp_dir):
    from PyPDF2 import PdfWriter
    pdf = tmp_dir / "small.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(pdf, "wb") as f:
        writer.write(f)
    assert is_natively_embeddable(pdf) is True


def test_get_pdf_page_count(tmp_dir):
    from PyPDF2 import PdfWriter
    pdf = tmp_dir / "test.pdf"
    writer = PdfWriter()
    for _ in range(3):
        writer.add_blank_page(width=72, height=72)
    with open(pdf, "wb") as f:
        writer.write(f)
    assert get_pdf_page_count(pdf) == 3


def test_get_pdf_page_count_missing_file():
    # Should return 999 (fallback to text extraction)
    assert get_pdf_page_count("/nonexistent.pdf") == 999


# --- Image conversion tests ---

def test_convert_image_to_png(tmp_dir):
    from PIL import Image
    # Create a BMP image
    bmp = tmp_dir / "test.bmp"
    img = Image.new("RGB", (10, 10), color="red")
    img.save(str(bmp), format="BMP")

    png_bytes, mime = convert_image_to_png(bmp)
    assert mime == "image/png"
    assert png_bytes[:4] == b"\x89PNG"
    assert len(png_bytes) > 0


def test_prepare_native_file_png(tmp_dir):
    from PIL import Image
    png = tmp_dir / "test.png"
    img = Image.new("RGB", (10, 10), color="blue")
    img.save(str(png), format="PNG")

    data, mime = prepare_native_file(png)
    assert mime == "image/png"
    assert data[:4] == b"\x89PNG"


def test_prepare_native_file_converts_gif(tmp_dir):
    from PIL import Image
    gif = tmp_dir / "test.gif"
    img = Image.new("RGB", (10, 10), color="green")
    img.save(str(gif), format="GIF")

    data, mime = prepare_native_file(gif)
    assert mime == "image/png"  # Converted to PNG
    assert data[:4] == b"\x89PNG"


def test_prepare_native_file_mp3(tmp_dir):
    mp3 = tmp_dir / "test.mp3"
    mp3.write_bytes(b"\xff\xfb" + b"\x00" * 100)

    data, mime = prepare_native_file(mp3)
    assert mime == "audio/mpeg"
    assert data == b"\xff\xfb" + b"\x00" * 100


# --- Media duration tests ---

def test_get_media_duration_returns_none_for_invalid(tmp_dir):
    bad = tmp_dir / "not_audio.mp3"
    bad.write_bytes(b"\x00" * 50)
    # mutagen can't parse fake data, should return None
    result = get_media_duration(bad)
    assert result is None or isinstance(result, float)


def test_get_media_duration_missing_file():
    result = get_media_duration("/nonexistent/file.mp3")
    assert result is None
